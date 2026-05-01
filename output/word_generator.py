from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row

def generate_word_doc(sector_name, sector_data, stories, datasets_info, output_path):
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Capstone Project Report")
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{sector_data['display']} Analytics System").font.size = Pt(18)

    doc.add_paragraph(f"Generated: {date.today()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Project Title
    add_heading(doc, "1. Project Title")
    doc.add_paragraph(f"Capstone Requirement Generator — {sector_data['display']} Multi-Entity Analytics Platform")

    # 2. Business Scenario
    add_heading(doc, "2. Business Scenario")
    doc.add_paragraph(
        f"GlobalEdge {sector_data['display']} Corporation (GEC) is a multinational organisation "
        f"operating across 12 countries with regional headquarters in Singapore. The organisation "
        f"processes over 90,000 records daily across its {sector_data['display'].lower()} operations. "
        f"The Chief Data Officer, Mr. Robert Harrington, has commissioned a full-stack data analytics "
        f"platform to unify all entity data, generate automated insights, and drive strategic decisions. "
        f"The platform spans Unix automation, MongoDB storage, Python pipelines, PySpark analytics, "
        f"Advanced SQL reporting, and Power BI dashboards."
    )

    # 3. Entity Schemas
    add_heading(doc, "3. Entity Schemas")
    for entity in sector_data["entities"]:
        add_heading(doc, f"3.{sector_data['entities'].index(entity)+1}  {entity['name']} Entity", level=2)
        doc.add_paragraph(f"Storage format: {entity['format'].upper()}")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Field Name", "Inferred Type", "Sample Range"
        for field in entity["fields"]:
            f = field.lower()
            if "id" in f:       dtype, sample = "VARCHAR(20)", "UUID or alphanumeric key"
            elif "date" in f:   dtype, sample = "DATE",        "2021-01-01 to 2024-12-31"
            elif "amount" in f or "price" in f or "cost" in f: dtype, sample = "DECIMAL(12,2)", "100.00 – 80,000.00"
            elif "age" in f:    dtype, sample = "INT",         "18 – 75"
            elif "pct" in f or "rate" in f: dtype, sample = "FLOAT", "0.1 – 99.9"
            elif "status" in f: dtype, sample = "VARCHAR(20)", "Active / Inactive / Pending"
            elif "name" in f:   dtype, sample = "VARCHAR(100)", "Non-Indian English names"
            else:               dtype, sample = "VARCHAR(50)",  "Categorical / free text"
            add_table_row(table, [field, dtype, sample])
        doc.add_paragraph()

    # 4. User Stories (all 8)
    add_heading(doc, "4. User Stories")
    for story_obj in stories:
        add_heading(doc, f"US-{story_obj['index']:02d} | {story_obj['module']} | Entity: {story_obj['entity']}", level=2)
        doc.add_paragraph(story_obj["story"])
        doc.add_paragraph()

    # 5. Dataset Summary
    add_heading(doc, "5. Dataset Summary")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Entity","Format","Rows","File Path"
    for d in datasets_info:
        add_table_row(table, [d["entity"], d["format"].upper(), f"{d['rows']:,}", d["path"]])

    # 6. Tech Module Summary
    add_heading(doc, "6. Technology Module Summary")
    modules_summary = [
        ("Unix Commands",    "15 story-based requirements using basic Unix commands only (no advanced tools). Each with input dataset and complete solution."),
        ("Shell Scripting",  "10 story-based requirements: 8 for analysis, 2 for data cleaning (Unix file handling). Complete shell scripts provided."),
        ("MongoDB",          "Cleaning + CRUD + Aggregation. 10+ story-based requirements. Final display queries included."),
        ("Python",           "File handling on minimum 3 datasets (80k–1L rows each). Cleaning pipelines with random sector selection."),
        ("PySpark Core",     "Dataset cleaning + MySQL insert statement generation. RDD: 20 stories. DataFrames: 20 stories. Spark SQL: 20 stories."),
        ("Advanced SQL",     "15 story-based requirements using ONLY stored functions and stored procedures. No simple SELECT queries."),
        ("Power BI",         "15 Power Query transformations + 10 DAX measures. Visualization dashboards using same datasets."),
        ("Analysis",         "PySpark combining multiple datasets. RDD 20 + Spark SQL 20 + DataFrames 20 analysis requirements."),
    ]
    for mod, desc in modules_summary:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{mod}: ")
        run.bold = True
        p.add_run(desc)

    # 7. AGILE Note
    add_heading(doc, "7. AGILE Compliance")
    doc.add_paragraph("AGILE sheet must be filled daily. File: output/agile_tracker.xlsx")
    doc.add_paragraph("Each user story must have: Story ID, Description, Acceptance Criteria, Priority (Must Have / Should Have / Could Have), and Status.")

    doc.save(output_path)
    print(f"  Word document saved: {output_path}")